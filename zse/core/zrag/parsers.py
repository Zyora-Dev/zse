"""
Document Parsers — Extract text from any format.

Supports: PDF, DOCX, HTML, TXT, CSV, JSON, Markdown
All parsers return a unified ParsedDocument with extracted text
and structural metadata.
"""

import csv
import io
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional


@dataclass
class ParsedSection:
    """A structural section from a parsed document."""

    text: str
    heading: Optional[str] = None
    level: int = 0  # heading level (1=h1, 2=h2, etc.)
    page: Optional[int] = None  # for PDFs
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Unified output from any parser."""

    text: str  # Full extracted text
    title: str = ""
    source_type: str = "unknown"
    sections: List[ParsedSection] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    original_size: int = 0


def parse_file(path: str) -> ParsedDocument:
    """Auto-detect format and parse a file."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = p.suffix.lower()
    raw = p.read_bytes()
    original_size = len(raw)

    parsers = {
        ".txt": _parse_text,
        ".md": _parse_markdown,
        ".json": _parse_json,
        ".csv": _parse_csv,
        ".html": _parse_html,
        ".htm": _parse_html,
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
    }

    parser = parsers.get(suffix)
    if parser is None:
        # Fallback: treat as plain text
        parser = _parse_text

    doc = parser(raw, p.name)
    doc.original_size = original_size
    doc.source_type = suffix.lstrip(".")
    if not doc.title:
        doc.title = p.stem
    return doc


def parse_bytes(data: bytes, filename: str) -> ParsedDocument:
    """Parse raw bytes with filename hint for format detection."""
    suffix = Path(filename).suffix.lower()
    parsers = {
        ".txt": _parse_text,
        ".md": _parse_markdown,
        ".json": _parse_json,
        ".csv": _parse_csv,
        ".html": _parse_html,
        ".htm": _parse_html,
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
    }
    parser = parsers.get(suffix, _parse_text)
    doc = parser(data, filename)
    doc.original_size = len(data)
    doc.source_type = suffix.lstrip(".") if suffix else "txt"
    if not doc.title:
        doc.title = Path(filename).stem
    return doc


# ---------------------------------------------------------------------------
# Individual Parsers
# ---------------------------------------------------------------------------


def _parse_text(data: bytes, filename: str) -> ParsedDocument:
    """Parse plain text."""
    text = data.decode("utf-8", errors="replace")
    return ParsedDocument(
        text=text,
        title=Path(filename).stem,
        sections=[ParsedSection(text=text)],
    )


def _parse_markdown(data: bytes, filename: str) -> ParsedDocument:
    """Parse Markdown — extract sections by headings."""
    text = data.decode("utf-8", errors="replace")
    sections: List[ParsedSection] = []
    title = Path(filename).stem

    # Split by headings
    heading_re = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)
    last_end = 0
    current_heading = None
    current_level = 0

    for m in heading_re.finditer(text):
        # Save previous section
        if last_end < m.start():
            section_text = text[last_end : m.start()].strip()
            if section_text:
                sections.append(
                    ParsedSection(
                        text=section_text,
                        heading=current_heading,
                        level=current_level,
                    )
                )

        current_level = len(m.group(1))
        current_heading = m.group(2).strip()
        if current_level == 1 and not title:
            title = current_heading
        last_end = m.end()

    # Last section
    if last_end < len(text):
        section_text = text[last_end:].strip()
        if section_text:
            sections.append(
                ParsedSection(
                    text=section_text,
                    heading=current_heading,
                    level=current_level,
                )
            )

    if not sections:
        sections = [ParsedSection(text=text)]

    return ParsedDocument(text=text, title=title, sections=sections)


def _parse_json(data: bytes, filename: str) -> ParsedDocument:
    """Parse JSON — flatten to readable text for LLM consumption."""
    text = data.decode("utf-8", errors="replace")
    try:
        obj = json.loads(text)
    except json.JSONDecodeError:
        return ParsedDocument(text=text, title=Path(filename).stem)

    lines: List[str] = []
    _flatten_json(obj, lines, depth=0)
    flat_text = "\n".join(lines)

    return ParsedDocument(
        text=flat_text,
        title=Path(filename).stem,
        sections=[ParsedSection(text=flat_text)],
        metadata={"original_type": type(obj).__name__},
    )


def _flatten_json(obj: Any, lines: List[str], depth: int, prefix: str = ""):
    """Recursively flatten JSON into readable lines."""
    indent = "  " * depth
    if isinstance(obj, dict):
        for key, val in obj.items():
            full_key = f"{prefix}.{key}" if prefix else key
            if isinstance(val, (dict, list)):
                lines.append(f"{indent}{key}:")
                _flatten_json(val, lines, depth + 1, full_key)
            else:
                lines.append(f"{indent}{key}: {val}")
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            if isinstance(item, (dict, list)):
                lines.append(f"{indent}[{i}]:")
                _flatten_json(item, lines, depth + 1, f"{prefix}[{i}]")
            else:
                lines.append(f"{indent}- {item}")
    else:
        lines.append(f"{indent}{obj}")


def _parse_csv(data: bytes, filename: str) -> ParsedDocument:
    """Parse CSV — convert to readable text with headers as context."""
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)

    if not rows:
        return ParsedDocument(text="", title=Path(filename).stem)

    headers = rows[0] if rows else []
    sections: List[ParsedSection] = []
    lines: List[str] = []

    # Header line
    if headers:
        lines.append("Columns: " + ", ".join(headers))
        lines.append("")

    # Data rows as key-value pairs
    for i, row in enumerate(rows[1:], 1):
        if headers:
            pairs = [f"{h}: {v}" for h, v in zip(headers, row) if v.strip()]
            lines.append(f"Row {i}: " + " | ".join(pairs))
        else:
            lines.append(" | ".join(row))

    flat_text = "\n".join(lines)
    sections.append(ParsedSection(text=flat_text))

    return ParsedDocument(
        text=flat_text,
        title=Path(filename).stem,
        sections=sections,
        metadata={"rows": len(rows) - 1, "columns": len(headers)},
    )


def _parse_html(data: bytes, filename: str) -> ParsedDocument:
    """Parse HTML — strip tags, extract text with structure."""
    text = data.decode("utf-8", errors="replace")

    # Try BeautifulSoup first
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(text, "html.parser")

        # Extract title
        title_tag = soup.find("title")
        title = title_tag.get_text().strip() if title_tag else Path(filename).stem

        # Remove script/style
        for tag in soup.find_all(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        # Extract sections by headings
        sections: List[ParsedSection] = []
        body = soup.find("body") or soup

        for heading in body.find_all(re.compile(r"^h[1-6]$")):
            level = int(heading.name[1])
            heading_text = heading.get_text().strip()

            # Collect text until next heading
            content_parts = []
            sibling = heading.find_next_sibling()
            while sibling and not re.match(r"^h[1-6]$", sibling.name or ""):
                t = sibling.get_text().strip()
                if t:
                    content_parts.append(t)
                sibling = sibling.find_next_sibling()

            if content_parts:
                sections.append(
                    ParsedSection(
                        text="\n".join(content_parts),
                        heading=heading_text,
                        level=level,
                    )
                )

        clean_text = soup.get_text(separator="\n")
        # Clean up whitespace
        clean_text = re.sub(r"\n{3,}", "\n\n", clean_text).strip()

        if not sections:
            sections = [ParsedSection(text=clean_text)]

        return ParsedDocument(text=clean_text, title=title, sections=sections)

    except ImportError:
        # Fallback: regex-based tag stripping
        clean = re.sub(r"<script[^>]*>.*?</script>", "", text, flags=re.DOTALL | re.IGNORECASE)
        clean = re.sub(r"<style[^>]*>.*?</style>", "", clean, flags=re.DOTALL | re.IGNORECASE)
        clean = re.sub(r"<[^>]+>", " ", clean)
        clean = re.sub(r"\s+", " ", clean).strip()
        return ParsedDocument(
            text=clean,
            title=Path(filename).stem,
            sections=[ParsedSection(text=clean)],
        )


def _parse_pdf(data: bytes, filename: str) -> ParsedDocument:
    """Parse PDF — extract text per page."""
    sections: List[ParsedSection] = []
    all_text_parts: List[str] = []
    title = Path(filename).stem

    # Try pypdf first
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))

        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            page_text = page_text.strip()
            if page_text:
                all_text_parts.append(page_text)
                sections.append(
                    ParsedSection(
                        text=page_text,
                        heading=f"Page {i + 1}",
                        page=i + 1,
                    )
                )

        # Try to get title from metadata
        meta = reader.metadata
        if meta and meta.title:
            title = meta.title

    except ImportError:
        # Try pdfplumber
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    page_text = page_text.strip()
                    if page_text:
                        all_text_parts.append(page_text)
                        sections.append(
                            ParsedSection(
                                text=page_text,
                                heading=f"Page {i + 1}",
                                page=i + 1,
                            )
                        )
        except ImportError:
            raise ImportError(
                "PDF parsing requires 'pypdf' or 'pdfplumber'. Install: pip install pypdf"
            )

    full_text = "\n\n".join(all_text_parts)
    return ParsedDocument(text=full_text, title=title, sections=sections)


def _parse_docx(data: bytes, filename: str) -> ParsedDocument:
    """Parse DOCX — extract text with paragraph structure."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("DOCX parsing requires 'python-docx'. Install: pip install python-docx")

    doc = DocxDocument(io.BytesIO(data))
    sections: List[ParsedSection] = []
    all_text_parts: List[str] = []
    title = Path(filename).stem

    current_heading = None
    current_level = 0
    current_parts: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower()

        if "heading" in style_name:
            # Flush previous section
            if current_parts:
                sections.append(
                    ParsedSection(
                        text="\n".join(current_parts),
                        heading=current_heading,
                        level=current_level,
                    )
                )
                current_parts = []

            # Parse heading level
            try:
                current_level = int(style_name.replace("heading", "").strip())
            except ValueError:
                current_level = 1
            current_heading = text
            if current_level == 1 and title == Path(filename).stem:
                title = text
        else:
            current_parts.append(text)

        all_text_parts.append(text)

    # Flush last section
    if current_parts:
        sections.append(
            ParsedSection(
                text="\n".join(current_parts),
                heading=current_heading,
                level=current_level,
            )
        )

    if not sections:
        full = "\n".join(all_text_parts)
        sections = [ParsedSection(text=full)]

    return ParsedDocument(
        text="\n".join(all_text_parts),
        title=title,
        sections=sections,
    )
